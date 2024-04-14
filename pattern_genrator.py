from random import randint, choice, shuffle
from docx import Document
from tqdm import tqdm

def main(doc, qid):
    number = 3#randint(3, 5)
    max_range = 99//number
    fn = randint(2, max_range)
    sum = fn
    num_list = [fn]
    try:
        for _ in range(1, number):
            if randint(0,1) == 1:
                rn = randint(1, min(max_range, sum))
            else:
                rn = (-1)*randint(1, min(max_range, sum))
            num_list.append(rn)
            sum = sum + rn
        doc.add_paragraph(f"[Q]")
        # test_case = "[Q]\n"
        # nums = '\n'.join(map(str, num_list))
        # print(''.join(map(str, num_list)))
        for num in num_list:
            doc.add_paragraph(f"{num}")
        # doc.add_paragraph(f"")
        doc.add_paragraph("[qtype] mcq")
        option = list(range(max(0, sum-10),sum+11))
        optt = []
        for _ in range(3):
            v = choice(option)
            option.remove(v)
            optt.append(v)
        if sum not in optt:
            optt.append(sum)
        else:
            optt.append(choice(option))
        shuffle(optt)
        for v, op in zip(optt, ["a", "b", "c", "d"]):
                if v == sum:
                    sum = op
                doc.add_paragraph(f"[{op}] {v}")

        doc.add_paragraph("[ans]")
        doc.add_paragraph(f"{sum}")
        doc.add_paragraph("[Marks] 2")
        doc.add_paragraph("+ve marks: 2 , -ve marks: 0")
        doc.add_paragraph(f"[sortid] {qid}")
        doc.add_paragraph()
            # test_case += f"{num}\n"
        # test_case += f"[ans]\n{sum}\n"
        # test_case += "[qtype] dtq\n"
        # test_case += "[Marks] 1\n"
        # test_case += "+ve marks: 1 , -ve marks: 0\n"
        # doc.add_paragraph(f"[sortid] {qid}")
        # test_case += f"[sortid] {qid}\n"
        # filename.write(test_case)
        # doc.close()
        # print(f"Q{qid} is done")
    except:
        main(doc, qid)

for id in tqdm(range(1, 81), total=80):
    doc = Document()
    # with        as f:
    for qid in range(1, 101):
        main(doc, qid)
    doc.save(f"t/Abacus-BL-{id}.docx")
        # f.close()